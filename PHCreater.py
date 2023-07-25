from copy import deepcopy
from io import StringIO
from dataclasses import dataclass, field
from abc import ABC, abstractmethod
from DirectSorter import DirectSorter
import docx
from docx.text.paragraph import Paragraph
from docx.document import Document
from docx.table import _Cell, Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
import docx
from docx2python import docx2python
import logging
logging.basicConfig(level=logging.DEBUG,
                    format="%(asctime)s %(levelname)s %(message)s",
                    datefmt="%Y-%m-%d %H:%M:%S")


def iter_block_items2(parent):
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            table = Table(child, parent)
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_block_items(cell)

def iter_block_items(parent):
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        yield child

    # for child in parent_elm.iterchildren():
    #     if isinstance(child, CT_P):
    #         yield Paragraph(child, parent)
    #     elif isinstance(child, CT_Tbl):
    #         table = Table(child, parent)
    #         for row in table.rows:
    #             for cell in row.cells:
    #                 yield from iter_block_items(cell)
@dataclass(slots=True)
class DocxCreater(ABC):
    directSorter: DirectSorter = field(default_factory=DirectSorter)

    def copyTestCasesToTestDoc(self):
        """ copies test cases from PH to test doc """
        with docx2python(self.directSorter.pathrunner.path + "/" + "t1.docx") as docx_content:
            logging.debug(f'docx2python data:{docx_content.document}')
        doc = docx.Document((self.directSorter.pathrunner.path + "/" + "t1.docx"))
        doc2 = docx.Document((self.directSorter.pathrunner.path + "/" + "t2.docx"))
        all_paras = doc.paragraphs
        # len(all_paras)
        # for para in all_paras:
        #     print(para.text)
        #     print("-------")
        # print("-----333333333333333333-----")
        # for table in doc.tables:
        #     print(table)
        #     print("-------")
        # template = doc.tables[0]
        # tbl = template._tbl
        # # Here we do the copy of the table
        # new_tbl = deepcopy(tbl)
        # print(new_tbl)
        # paragraph = doc2.add_paragraph()
        # # After that, we add the previously copied table
        # paragraph._p.addnext(new_tbl)

        #
        # for i in iter_block_items2(doc):
        #     print(i.text)
        for para in all_paras:
            print()
            logging.debug(f'para data:{type(para)} {para}')

        for block in iter_block_items(doc):
            if isinstance(block, CT_P):
                Paragraph(block, doc)

            elif isinstance(block, CT_Tbl):
                """
                we need to find a way to loop through the table and find the key words 
                if found we copy the whole table to the new document 
                """
                table = Table(block, doc)
                new_tbl = deepcopy(block)
                for cells in table.row_cells:
                    """ todo how to get the correct position of table to copy """

                logging.debug(f'new_tbl data:{table.row_cells}')
                paragraph = doc2.add_paragraph()
                # After that, we add the previously copied table
                paragraph._p.addnext(new_tbl)

        doc2.save(self.directSorter.pathrunner.path + "/" + "t3.docx")